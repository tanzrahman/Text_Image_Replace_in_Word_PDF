import datetime

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt


def replace_word_with_image(doc_path, sign, date, image_path, output_path):
    doc = Document(doc_path)
    # Iterate through tables
    signed = False
    date_added = False
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    print(paragraph.text)

                    if date in paragraph.text:
                        date_text = "\t"+str(datetime.date.today())
                        paragraph.text = paragraph.text.replace(date,date_text)
                        date_added = True

                    if sign in paragraph.text:
                        paragraph.text = paragraph.text.replace(sign,"")

                        # Add image to the cell
                        run = paragraph.add_run()
                        run.add_picture(image_path, width=Inches(1.5))  # Adjust size as needed
                        run.underline = True

                        doc.save(output_path)
                        signed = True
        if(signed and date_added):
            doc.save(output_path)
            return


# Example usage
replace_word_with_image("E:/input.docx", '[SD_SIGNATURE]', '[SIGNING_DATE]', "E:/sign.png", "E:/output.docx")

